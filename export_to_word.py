import os
import sys

# Try to import python-docx
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Error: 'python-docx' library is not installed.")
    print("Please install it by running: pip install python-docx")
    sys.exit(1)

def is_binary(file_path):
    """
    Simple check to see if a file is binary.
    Reads a small chunk and checks for null bytes.
    """
    try:
        with open(file_path, 'rb') as f:
            chunk = f.read(1024)
            if b'\0' in chunk:
                return True
    except Exception:
        return True
    return False

def collect_files(root_dir, skip_dirs, skip_exts):
    """
    Walks the directory and yields file paths that match criteria.
    """
    for dirpath, dirnames, filenames in os.walk(root_dir):
        # Modify dirnames in-place to skip directories
        dirnames[:] = [d for d in dirnames if d not in skip_dirs]
        
        for filename in filenames:
            ext = os.path.splitext(filename)[1].lower()
            if ext in skip_exts:
                continue
            
            yield os.path.join(dirpath, filename)

def create_word_doc(root_dir, output_file):
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(10)

    # Directories to ignore (artifacts, git, etc.)
    SKIP_DIRS = {
        '.git', '.vs', 'build', 'out', 'bin', 'obj', 
        '__pycache__', '.vscode', 'node_modules', 'artifacts',
        '.gemini' # Skip agent files
    }
    
    # Extensions to ignore (binaries, images, etc.)
    SKIP_EXTS = {
        '.exe', '.dll', '.obj', '.o', '.bin', '.img', '.iso', 
        '.pdb', '.idb', '.ilk', '.suo', '.user', '.db', 
        '.bmp', '.png', '.jpg', '.jpeg', '.gif', '.ico',
        '.pyc', '.zip', '.tar', '.gz', '.7z', '.pdf', '.docx'
    }

    print(f"Scanning directory: {root_dir}")
    print("Generating Word document...")

    count = 0
    for file_path in collect_files(root_dir, SKIP_DIRS, SKIP_EXTS):
        # Check for binary content before processing
        if is_binary(file_path):
            print(f"Skipping binary file: {os.path.relpath(file_path, root_dir)}")
            continue

        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                
            rel_path = os.path.relpath(file_path, root_dir)
            
            # Add File Name Header
            doc.add_heading(rel_path, level=1)
            
            # Add File Content
            # Using a simplified approach: just adding text. 
            # For code-block like styling in Word, we use a different font or style, 
            # but 'Consolas' on Normal style covers it reasonably well for a dump.
            p = doc.add_paragraph(content)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Add a page break or spacing
            doc.add_page_break()
            
            count += 1
            if count % 10 == 0:
                print(f"Processed {count} files...")
                
        except Exception as e:
            print(f"Error reading {file_path}: {e}")

    print(f"Saving to {output_file}...")
    doc.save(output_file)
    print(f"Done! Successfully exported {count} files.")

if __name__ == "__main__":
    # Current workspace directory
    workspace_dir = os.getcwd() 
    output_filename = "Codebase_Export.docx"
    
    create_word_doc(workspace_dir, output_filename)
