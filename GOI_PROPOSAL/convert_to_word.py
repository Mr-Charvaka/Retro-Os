import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(output_filename, input_dir):
    doc = Document()
    
    # Correct Word style for a document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    files = [
        "Part_1_Executive_Summary_and_Strategic_Vision.md",
        "Part_2_Technical_Architecture_and_Kernel_Design.md",
        "Part_3_Secure_Networking_and_TLS_Stack.md",
        "Part_4_AI_and_Future_Roadmap.md",
        "Part_5_Impact_and_Conclusion.md",
        "Part_6_Technical_API_Reference.md",
        "Part_7_Performance_Benchmarks.md",
        "Part_8_Roadmap_and_Conclusion.md"
    ]

    for file_name in files:
        file_path = os.path.join(input_dir, file_name)
        if not os.path.exists(file_path):
            print(f"Skipping {file_name}: File not found.")
            continue
            
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith("# "):
                doc.add_heading(line[2:], level=0)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=2)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=3)
            elif line.startswith("---"):
                doc.add_page_break()
            elif line.startswith("|"): # Table - basic handling
                # Since professional tables are complex, let's just add it as a paragraph for now
                doc.add_paragraph(line)
            elif line.startswith("* "): # Bullet
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith("1. "): # Numbered
                doc.add_paragraph(line[3:], style='List Number')
            else:
                doc.add_paragraph(line)
                
        # Add a page break between volumes
        doc.add_page_break()

    doc.save(output_filename)
    print(f"Successfully saved to {output_filename}")

if __name__ == "__main__":
    input_directory = r"d:\Retro-Os-MATTERS HTTPS ACCESSIBLE\GOI_PROPOSAL"
    output_path = r"d:\Retro-Os-MATTERS HTTPS ACCESSIBLE\Detailed Report.docx"
    markdown_to_docx(output_path, input_directory)
